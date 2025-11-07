from flask import Flask, render_template, redirect, request, session, flash, url_for, send_file,send_from_directory, jsonify, abort
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import secrets
import json
import time
import tempfile
import whisper
import whisperx
import pandas as pd
import openai
import requests
from openai import OpenAI
from werkzeug.utils import secure_filename
from pydub import AudioSegment 
from pydub.silence import detect_nonsilent
from flask_socketio import SocketIO, emit
from io import BytesIO
from faster_whisper import WhisperModel
from transcriber import register_routes
from datetime import datetime
from urllib.parse import quote
from flask_cors import CORS
import tiktoken
import zipfile


app = Flask(__name__)
CORS(app)# 允許跨域請求
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 限制上傳檔案大小 100MB
socketio = SocketIO(app, cors_allowed_origins="*")# 初始化 WebSocket
register_routes(app, socketio) # 導入 transcriber.py 中的路由
app.secret_key = 'your_secret_key'
client = OpenAI(api_key="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")
OLLAMA_API_URL = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "qwen:7b-chat"  #


#==================設定專案中使用的資料夾、支援檔案格式、自動建立資料夾（若不存在）===============
USER_FILE = "users.txt"  # 使用者資料存儲檔案

RECORDINGS_FOLDER = 'recordings'
UPLOAD_FOLDER = 'uploads'
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'outputs')
#OUTPUT_FOLDER = 'outputs'
TRANSCRIBED_FILES_DIR = "transcribed_files"
ALLOWED_EXTENSIONS = {'webm','mp3', 'wav', 'm4a'}

os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TRANSCRIBED_FILES_DIR, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

# 設定 WhisperX
device = "cpu"  
batch_size = 16
compute_type = "int8"

# 載入 WhisperX 模型
model = whisperx.load_model("base", device, compute_type=compute_type)
diarize_model = whisperx.DiarizationPipeline(device=device, use_auth_token="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")

# 儲存使用者到記事本
def save_user_to_file(username, password, api_key=None):
    with open(USER_FILE, "a", encoding="utf-8") as file:
        if api_key:
            file.write(f"{username},{password},{api_key}\n")
        else:
            file.write(f"{username},{password}\n")

# 從記事本讀取使用者資料
def load_users():
    users = {}
    if os.path.exists(USER_FILE):
        with open(USER_FILE, "r", encoding="utf-8") as file:
            for line in file:
                parts = line.strip().split(",")
                if len(parts) == 2:
                    username, password = parts
                    users[username] = {"password": password, "api_key": None}
                elif len(parts) == 3:
                    username, password, api_key = parts
                    users[username] = {"password": password, "api_key": api_key}
    return users


# 註冊頁面
@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        users = load_users()

        if username in users:
            flash("使用者名稱已存在！", "danger")
        else:
            save_user_to_file(username, password)
            flash("註冊成功！請登入", "success")
            return redirect("/login")

    return render_template("register.html")

# 登入頁面
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        users = load_users()

        if username not in users:
            flash("您未註冊帳號！", "danger")  # 若帳號不存在
        elif users[username]["password"] == password:
            session["username"] = username
            return redirect("/apikey")  # 登入成功後跳轉到首頁
        else:
            flash("帳號或密碼錯誤！", "danger")

    return render_template("login.html")

# 首頁（登入後顯示）
@app.route("/")
def dashboard():
    if 'username' not in session:
        return redirect(url_for('login'))  # 如果未登入，跳回登入頁面

    return render_template('apikey.html')

@app.route("/apikey")
def apikey():
    if 'username' not in session:
        flash("請先登入！", "danger")  # 提示未登入
        return redirect(url_for('login'))  # 如果未登入，跳回登入頁面

    return render_template("apikey.html")  # 顯示 API Key 頁面

def read_api_key(username):
    if not os.path.exists(USER_FILE):
        return "尚未產生 API Key"
    
    with open(USER_FILE, "r") as file:
        for line in file:
            data = line.strip().split(",")  # 根據「,」拆分
            if len(data) == 3:  # 確保有 username, password, apikey
                user, password, key = data
                if user == username:
                    return key  # 回傳 API Key

    return "尚未產生 API Key"

# 生成新的 API Key
@app.route('/generate_apikey', methods=['POST'])
def generate_apikey():
    if 'username' not in session:
        return "請先登入", 401

    username = session['username']
    new_key = secrets.token_hex(8)  # 生成 16 字節的 API Key

    # 讀取所有使用者 API Key
    users = []
    if os.path.exists(USER_FILE):
        with open(USER_FILE, "r") as file:
            for line in file:
                data = line.strip().split(",")
                if len(data) == 3:
                    users.append(data)

    # 更新 API Key
    for user in users:
        if user[0] == username:
            user[2] = new_key  # 更新 API Key

    # 寫回 `users.txt`
    with open(USER_FILE, "w") as file:
        for user in users:
            file.write(",".join(user) + "\n")

    return new_key

# 取得目前的 API Key
@app.route('/get_apikey', methods=['GET'])
def get_apikey():
    if 'username' not in session:
        return "請先登入", 401

    username = session['username']
    return read_api_key(username)

# 檢查文件是否為允許的格式
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# 使用 Whisper 轉錄音檔
def transcribe_audio(audio_path, language, model_name):
    model = whisper.load_model(model_name)
    result = model.transcribe(audio_path, language=language)
    return result

def split_text_by_tokens(text, max_tokens=3000, model_name="gpt-3.5-turbo"):
    encoding = tiktoken.encoding_for_model(model_name)
    tokens = encoding.encode(text)
    chunks = []
    start = 0

    while start < len(tokens):
        end = start + max_tokens
        chunk = encoding.decode(tokens[start:end])
        chunks.append(chunk)
        start = end

    return chunks

def split_large_text(text, chunk_size=5000):
    return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

# 音訊分段轉文字給 Dify 做校正
def send_to_dify(transcript_text):
    url = 'https://api.dify.ai/v1/chat-messages'
    headers = {
        'Authorization': 'Bearer app-ga2ISeylxzU4CU7lwYc3brej',
        'Content-Type': 'application/json'
    }

    transcript_parts = split_text_by_tokens(transcript_text, max_tokens=3000)
    all_responses = []

    for i, part in enumerate(transcript_parts, 1):
        data = {
            'inputs': {},
            'query': f"【第 {i} 段校正】請幫我校正這段逐字稿，並修正錯字，保證文字要是繁體中文：{part}",
            'response_mode': 'blocking',
            'user': f'test_user_{i:03d}'
        }

        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            result = response.json()
            all_responses.append(f"【第 {i} 段校正結果】 {result['answer']}")
        else:
            all_responses.append(f"❌ 第 {i} 段處理失敗：{response.status_code} {response.text}")

    return "\n\n".join(all_responses)
    
def create_zip_file(files):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for filename, file_content in files.items():
            zip_file.writestr(filename, file_content)
    zip_buffer.seek(0)
    return zip_buffer


#======================音檔轉文字，上傳音訊，進行 Whisper 轉錄=========================
@app.route('/voicetotext', methods=['GET', 'POST'])
def index():
    transcript = None  # 存放轉錄結果
    output_filename = None  # 儲存轉錄檔案名稱
    error_message = None  # 用於錯誤提示
    transcription_time = None  # 轉錄所需時間
    dify_response = None

    if request.method == 'POST':
        if 'audio_file' not in request.files or request.files['audio_file'].filename == "":
            error_message = "請選擇音訊檔案"
        else:
            # 取得表單數據
            audio_file = request.files['audio_file']

            ALLOWED_EXTENSIONS = {'mp3', 'wav', 'm4a'}
            def allowed_file(filename):
                return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

            # 檢查音訊格式
        if not allowed_file(audio_file.filename):
                error_message = "不支援的音訊格式，請上傳 mp3、m4a 或 wav 檔案"
        else:

            language = request.form.get('language')  # 預設語言為英文
            model_name = request.form.get('model')  # 預設 Whisper base 模型
            output_format = request.form.get('output_format')  # 預設輸出格式為 txt

            # 儲存音訊檔案
            audio_path = os.path.join(UPLOAD_FOLDER, audio_file.filename)
            audio_file.save(audio_path)

            # 轉錄音訊並計算時間
            translate = request.form.get('translate') == 'on'
            # 🔴 檢查語言為英文時是否選了翻譯
            if translate and language == 'en':
                error_message = "英文音訊無法翻譯為英文，請取消翻譯選項。"
                return render_template('index.html', transcript=None, error=error_message, output_file=None, transcription_time=None)
            start_time = time.time()  # 開始計時
            transcript, segments = process_audio(audio_path, language, model_name, translate)
            dify_response = send_to_dify(transcript)
            end_time = time.time()  # 結束計時

            transcription_time = end_time - start_time  # 計算所需時間

            if transcript:
                # 生成對應格式的輸出檔案
                output_filename = create_output_file(audio_file.filename, transcript, segments, output_format)
            else:
                error_message = "轉錄失敗，請檢查音檔"
        print(dify_response)
    return render_template('index.html', transcript=transcript, dify_reply=dify_response ,error=error_message, output_file=output_filename, transcription_time=transcription_time)

def process_audio(audio_path, language, model_name, translate=False):
    try:
        model = whisper.load_model(model_name)
        result = model.transcribe(audio_path, language=language, task="translate" if translate else "transcribe")
        return result["text"], result["segments"]
    except Exception as e:
        print(f"轉錄時發生錯誤: {e}")
        return None, None
    

def create_output_file(filename, transcript, segments, output_format):
    """根據用戶選擇的格式生成對應的檔案"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"{os.path.splitext(filename)[0]}_{timestamp}.{output_format}"
    #output_filename = f"{os.path.splitext(filename)[0]}.{output_format}"
    output_path = os.path.join(OUTPUT_FOLDER, output_filename)

    if output_format == "txt":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(transcript)

    elif output_format == "srt":
        with open(output_path, "w", encoding="utf-8") as f:
            for i, segment in enumerate(segments, start=1):
                start_time = format_timestamp(segment["start"])
                end_time = format_timestamp(segment["end"])
                f.write(f"{i}\n{start_time} --> {end_time}\n{segment['text']}\n\n")

    elif output_format == "tsv":
        with open(output_path, "w", encoding="utf-8") as f:
            for segment in segments:
                f.write(f"{segment['start']}\t{segment['end']}\t{segment['text']}\n")
    print(f"寫入檔案成功：{output_path}")
    return output_filename  # 只回傳檔名，不回傳完整路徑

def format_timestamp(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    seconds = int(seconds % 60)
    milliseconds = int((seconds % 1) * 1000)
    return f"{hours:02}:{minutes:02}:{seconds:02},{milliseconds:03}"

@app.route('/download/<filename>')
def download_file(filename):
    file_path = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(file_path):
        return "檔案不存在", 404
    response = send_file(file_path, as_attachment=True, mimetype='application/octet-stream')
    response.headers["Content-Disposition"] = f"attachment; filename*=UTF-8''{quote(filename)}"
    return response

def transcribe_audio(file_path):
    audio = whisperx.load_audio(file_path)
    result = model.transcribe(audio)
    return pd.DataFrame(result["segments"])

# 摘要會議內容
def generate_meeting_summary(transcription_text):
    # 使用 OpenAI 進行會議摘要
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "請幫我做會議摘要。"},
            {"role": "user", "content": transcription_text}
        ]
    )
    summary = response.choices[0].message['content'].strip()
    return summary

#==================會議記錄轉文字=============================
@app.route("/meeting", methods=["GET", "POST"])
def meeting():
    table_html = None
    output_file = None
    transcription_time = None  # 新增變數來儲存轉錄時間
    error_message = None  # 儲存錯誤訊息
    meeting_summary = None
    todo_list = None  # ✅ 新增代辦清單變數



    if request.method == "POST":
        if "audio_file" not in request.files or request.files["audio_file"].filename == "":
            error_message = "請選擇音訊檔案"
        else:
            file = request.files["audio_file"]

            # 允許的音訊格式
            ALLOWED_EXTENSIONS = {"mp3", "wav", "m4a"}
            def allowed_file(filename):
                return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

            # 檢查音訊格式
            if not allowed_file(file.filename):
                error_message = "不支援的音訊格式，請上傳 mp3、m4a 或 wav 檔案"
            else:
                # 儲存檔案
                file_path = os.path.join(UPLOAD_FOLDER, file.filename)
                file.save(file_path)
                
            # 記錄轉錄開始時間
            start_time = time.time()

            # 進行轉錄
            audio = whisperx.load_audio(file_path)
            result = model.transcribe(audio, batch_size=batch_size)
            transcription_df = pd.DataFrame(result["segments"])

            # 說話者辨識
            diarize_segments = diarize_model(file_path)
            speaker_df = pd.DataFrame(diarize_segments)

            # 初始化 speaker 欄位
            transcription_df["speaker"] = "Unknown"

            time_tolerance = 1.0
            for i, row in transcription_df.iterrows():
                best_match = None
                best_overlap = 0

                for _, speaker_row in speaker_df.iterrows():
                    speaker_start = speaker_row["start"] - time_tolerance
                    speaker_end = speaker_row["end"] + time_tolerance
                    overlap = min(row["end"], speaker_end) - max(row["start"], speaker_start)

                    if overlap > best_overlap:
                        best_overlap = overlap
                        best_match = speaker_row["speaker"]

                if best_match:
                    transcription_df.at[i, "speaker"] = best_match

                transcription_df = transcription_df[['speaker', 'start', 'end' ,'text']]

            # 下載檔案資訊:說話者與文字
            text_output = "\n".join([f"{row['speaker']}:{row['start']}-{row['end']}:{row['text']}" for _, row in transcription_df.iterrows()])


                # 用 Ollama 分析摘要
            try:
                all_text = " ".join(transcription_df["text"].tolist())
                # 強調要求返回中文摘要
                ollama_prompt = f"""
                     你是一個擅長整理會議紀錄的【繁體中文】專業助理，請根據以下內容整理成簡潔且自然流暢的繁體中文會議摘要：

                    會議內容如下：
                    {all_text}

                    請遵守以下要求：
                    - 請**只使用繁體中文**回答
                   
                    - 摘要要以簡單、清楚的方式呈現，語氣自然流暢，避免過於正式的條列式
                    - 保留重要的英文專有名詞（如 AI、API、Python 等），但整體語言維持繁體中文
                    - **不要**使用英文進行說明或總結

                    請注意：**不論輸入為何，請始終要以繁體中文作答**
                    """

                response = requests.post(OLLAMA_API_URL, json={
                    "model": OLLAMA_MODEL,
                    "prompt": ollama_prompt,
                    "stream": False,
                    "language": "zh",  # 確保語言設置為中文
                    "temperature": 0.3
                }, timeout=60)

                if response.status_code == 200:
                    meeting_summary = response.json().get("response", "").strip()
                else:
                    meeting_summary = f"⚠️ Ollama 摘要失敗：{response.status_code} {response.text}"
            except Exception as e:
                meeting_summary = f"⚠️ Ollama 摘要失敗：{str(e)}"

            
            # ✅ 使用摘要產出代辦清單
            if all_text:
                    try:
                        todo_prompt = f"""
              你是一位擅長理解會議內容並整理行動項目的【繁體中文】專業助理。請根據以下的「逐字稿內容」，整理出所有需要執行的繁體中文代辦事項（To-do List），僅使用**繁體中文作答**。

              📝 **以下是會議逐內容**：
                {all_text}

                請根據以下要求輸出：

                📌 **輸出格式規則**：
                -  一定要使用**繁體中文**
                - 每一項代辦請以「-」開頭條列
                - 每項代辦應為**具體可執行的任務**（例如：完成某報告、安排某會議、修改某功能、確認某事…）
                - 盡可能根據內容，為代辦事項進行簡單分類（例如：工程項目、行政事項、行銷任務、客戶追蹤、資料確認等），可加上分類名稱作為前綴（例如：「【工程】- 修正登入錯誤問題」）
                - 回覆請僅使用**繁體中文**，請勿穿插英文說明
                - **不要**使用英文進行說明或總結

                請注意：**不論輸入為何，請始終要以繁體中文作答**

                
                """
                        response_todo = requests.post(OLLAMA_API_URL, json={
                            "model": OLLAMA_MODEL,
                            "prompt": todo_prompt,
                            "stream": False,
                            "language": "zh",
                            "temperature": 0.3
                        }, timeout=60)

                        if response_todo.status_code == 200:
                            todo_list = response_todo.json().get("response", "").strip()
                        else:
                            todo_list = f"⚠️ Ollama 代辦產出失敗：{response_todo.status_code} {response_todo.text}"
                    except Exception as e:
                        todo_list = f"⚠️ Ollama 代辦產出失敗：{str(e)}"

                          
            # 存成 txt 檔
            output_file = f"{file.filename}.txt"  # 儲存檔案的相對路徑
            output_path = os.path.join(UPLOAD_FOLDER, output_file)  # 取得正確路徑
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text_output)

                if meeting_summary:
                    f.write("\n\n=== 會議摘要 ===\n")
                    f.write(meeting_summary)
                if todo_list:
                    f.write("\n\n=== 代辦事項 ===\n")
                    f.write(todo_list)

            # 轉成 HTML 表格顯示在前端
            table_html = transcription_df.to_html(classes="table table-bordered", index=False)

            # 記錄轉錄結束時間
            end_time = time.time()
            transcription_time = round(end_time - start_time, 2)  # 計算並四捨五入到小數點後兩位

    return render_template("meeting.html", 
                           transcription=table_html, 
                           output_file=output_file, 
                           transcription_time=transcription_time,
                           meeting_summary=meeting_summary,
                           todo_list=todo_list)


@app.route('/download_meeting/<filename>', methods=['GET'])
def download_meeting(filename):
    # 確保這裡的路徑是正確的
    file_path = os.path.join(UPLOAD_FOLDER, filename)  # 使用正確的 uploads 路徑
    file_path = os.path.normpath(file_path)  # 標準化路徑

    print(f"下載檔案的路徑: {file_path}")
    
    if os.path.exists(file_path):
        return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)
    else:
        flash("檔案不存在！", "danger")
        return redirect(url_for('meeting'))

# 下載 Word 版本會議紀錄
@app.route('/download_meeting_word/<filename>', methods=['GET'])
def download_meeting_word(filename):
    file_path = os.path.join(UPLOAD_FOLDER, filename)  # 原始 txt 檔案路徑
    word_filename = filename.replace(".txt", ".docx")  # Word 檔案名稱
    word_path = os.path.join(UPLOAD_FOLDER, word_filename)  # Word 檔案儲存路徑

    if os.path.exists(file_path):
        # 讀取 txt 檔案內容
        with open(file_path, "r", encoding="utf-8") as f:
            text_content = f.read()

        # 轉換為表格格式的內容
        table_data = []
        for line in text_content.split("\n"):
            if line.strip():  # 排除空行
                parts = line.split(":")
                if len(parts) >= 3:
                    speaker, time_range, text = parts[0], parts[1], ":".join(parts[2:])
                    table_data.append([speaker, time_range, text])

        # 建立 Word 文件
        doc = Document()
        doc.add_heading('會議記錄', 0)

        # 建立表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '說話者'
        hdr_cells[1].text = '時間範圍'
        hdr_cells[2].text = '內容'

        # 填入資料
        for row_data in table_data:
            row_cells = table.add_row().cells
            for i, data in enumerate(row_data):
                row_cells[i].text = data
                
        # ✅ 加入會議摘要（如果存在）
        if "=== 會議摘要 ===" in text_content:
            summary = text_content.split("=== 會議摘要 ===")[-1].strip()
            if summary:
                doc.add_paragraph("\n=== 會議摘要 ===", style='Heading 2')
                doc.add_paragraph(summary)

        # ✅ 加入代辦事項（如果存在）
        if "=== 代辦事項 ===" in text_content:
            todo_section = text_content.split("=== 代辦事項 ===")[-1].strip()
            if todo_section:
                doc.add_paragraph("\n=== 代辦事項 ===", style='Heading 2')
                doc.add_paragraph(todo_section)
        
        # 儲存 Word 檔案
        doc.save(word_path)

        # 提供 Word 檔案下載
        return send_file(word_path, as_attachment=True)
    else:
        flash("檔案不存在！", "danger")
        return redirect(url_for('meeting'))


# 儲存轉錄結果
def save_transcription(result, filename, output_format):
    output_path = os.path.join("transcriptions", f"{filename}.{output_format}")
    transcription_df = pd.DataFrame(result["segments"])
    
    if output_format == "csv":
        transcription_df.to_csv(output_path, index=False)
    elif output_format == "txt":
        with open(output_path, "w") as f:
            for segment in result["segments"]:
                f.write(f"{segment['start']} - {segment['end']} : {segment['text']}\n")

    # 加入輸出檔案路徑的 print
    print(f"儲存轉錄檔案的路徑: {output_path}")

    return output_path


#===================及時錄音轉文字================
@app.route("/transcribe", methods=["GET", "POST"])
def transcribe():
    if request.method == "GET":
        return render_template("transcribe.html")

    if 'audio_file' not in request.files:
        return jsonify({"error": "未找到音檔"}), 400

    file = request.files['audio_file']

    if file.filename == '':
        return jsonify({"error": "請選擇一個音檔"}), 400

    filename = secure_filename(file.filename)

    if allowed_file(filename):
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # 檢查音檔是否為 webm 格式，並將其轉換為 wav 格式
        if filename.endswith('.webm'):
            audio = AudioSegment.from_file(filepath, format="webm")
            filename = filename.rsplit('.', 1)[0] + '.wav'
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            audio.export(filepath, format="wav")

        # 使用 Whisper 進行轉錄
        transcription = transcribe_audio_with_whisper(filepath)

        return render_template("transcribe.html", transcription=transcription)

    return jsonify({"error": "不支援的檔案格式"}), 400

def transcribe_audio_with_whisper(filepath):
    model = whisper.load_model("base")  # 載入 Whisper 模型
    result = model.transcribe(filepath)  # 執行轉錄
    return result["text"]  # 返回轉錄的文字內容



@app.route('/latest_file', methods=['GET'])
def get_latest_file():
    # 目標資料夾
    folder = 'outputs'
    
    # 取得資料夾中的所有檔案
    files = os.listdir(folder)
    # 排除不是檔案的項目
    files = [f for f in files if os.path.isfile(os.path.join(folder, f))]
    
    if not files:
        return jsonify({"error": "No files found"}), 404
    
    # 根據檔案的修改時間來找出最新的檔案
    latest_file = max(files, key=lambda f: os.path.getmtime(os.path.join(folder, f)))
    
    return jsonify({"filename": latest_file})


# 上傳音檔並轉錄，無需表單，直接在 /upload 完成
@app.route("/upload", methods=["POST"])
def upload():
    # 檢查音檔是否存在
    file = request.files.get('audio_file')  # 從 POST 請求中提取音檔
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # 取得語言與模型選項，從 POST 請求參數中提取
        language = request.form.get('language', 'en')  # 默認為英文
        model_name = request.form.get('model', 'base')  # 默認使用 base 模型

        # 進行轉錄
        transcription = transcribe_audio(filepath, language, model_name)

        # 獲取所選的檔案格式
        output_format = request.form.get('output_format', 'txt')  # 默認格式為 txt
        output_file = save_transcription(transcription, filename.rsplit('.', 1)[0], output_format)

        # 返回轉錄檔案的下載鏈接
        return send_file(output_file, as_attachment=True)
    
    else:
        return jsonify({"error": "請上傳有效的音檔格式！"}), 400  # 檢查音檔格式

@app.route("/change-password", methods=["GET", "POST"])
def change_password():
    if 'username' not in session:
        return redirect('/login')  # 如果用戶沒有登入，跳轉到登入頁面
    
    username = session['username']
    users = load_users()

    if request.method == "POST":
        old_password = request.form["old_password"]
        new_password = request.form["new_password"]

        # 檢查舊密碼是否正確
        if username in users and users[username]['password'] == old_password:
            # 更新密碼
            users[username]['password'] = new_password
            # 更新檔案中的資料
            with open(USER_FILE, "w", encoding="utf-8") as file:
                for user, info in users.items():
                    if 'api_key' in info:
                        file.write(f"{user},{info['password']},{info['api_key']}\n")
                    else:
                        file.write(f"{user},{info['password']}\n")
            flash("密碼更改成功！", "success")
            return redirect("/home")  # 密碼更改成功後跳轉
            
        else:
            flash("舊密碼錯誤！", "danger")  # 如果舊密碼錯誤，顯示錯誤訊息

    return render_template("change_password.html")


# 登出
@app.route("/logout")
def logout():
    session.pop("username", None)
    flash("已登出", "info")
    return redirect("/login")

if __name__ == "__main__":
    if not os.path.exists(USER_FILE):
        open(USER_FILE, "w").close()  # 若無檔案則建立
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)
    port = int(os.environ.get("PORT", 8080))
    socketio.run(app, debug=True, host='0.0.0.0', port=port)
